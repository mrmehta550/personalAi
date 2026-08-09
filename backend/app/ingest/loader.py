import os
from typing import List
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
from app.core.logger import logger

class DocumentLoader:
    @staticmethod
    def load_file(file_path: str, filename: str) -> List[Document]:
        ext = os.path.splitext(filename)[1].lower()
        docs = []

        try:
            if ext == ".pdf":
                reader = PdfReader(file_path)
                text = ""
                for page_idx, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        docs.append(Document(
                            page_content=page_text,
                            metadata={"source_file": filename, "page": page_idx + 1}
                        ))
            elif ext == ".docx":
                doc = DocxDocument(file_path)
                full_text = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text)
                docs.append(Document(
                    page_content="\n".join(full_text),
                    metadata={"source_file": filename}
                ))
            elif ext in [".txt", ".md"]:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                docs.append(Document(
                    page_content=text,
                    metadata={"source_file": filename}
                ))
            else:
                raise ValueError(f"Unsupported file format: {ext}")

        except Exception as e:
            logger.error(f"Error loading document {filename}: {e}")
            raise e

        return docs
